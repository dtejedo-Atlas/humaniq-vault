"""
Document Parser v2.1
====================
Extractor de texto de documentos PDF y DOCX.

NUEVO en v2.1:
- Soporte para PDFs de múltiples columnas (CVs de Canva, plantillas con diseño)
- Detección automática de layout multi-columna basada en coordenadas x
- Ordenamiento inteligente del texto para preservar orden lógico de lectura

Incluye fallback automático de OCR para PDFs escaneados (imágenes).

REQUISITOS DEL SISTEMA:
-----------------------
Para el OCR de PDFs escaneados, el servidor necesita tener instalados:
- tesseract-ocr: Motor de OCR
- tesseract-ocr-spa: Datos de idioma español
- poppler-utils: Herramientas para convertir PDF a imagen (pdftoppm)

Instalación en Debian/Ubuntu:
    apt-get install tesseract-ocr tesseract-ocr-spa poppler-utils

Instalación en macOS:
    brew install tesseract poppler
"""

import pdfplumber
from docx import Document
from pathlib import Path
import io
import logging
from text_utils import clean_text_encoding
from collections import defaultdict

# Configurar logger
logger = logging.getLogger(__name__)

# Umbral mínimo de caracteres para considerar que un PDF tiene texto extraíble
MIN_TEXT_THRESHOLD = 20

# Umbral para detectar columnas (porcentaje del ancho de página)
COLUMN_GAP_THRESHOLD = 0.15  # 15% del ancho = gap entre columnas


def _detect_columns(words: list, page_width: float) -> list:
    """
    Detecta columnas en una página basándose en las coordenadas x de las palabras.
    
    Args:
        words: Lista de diccionarios con 'x0', 'x1', 'top', 'bottom', 'text'
        page_width: Ancho de la página
        
    Returns:
        Lista de rangos de columnas [(x_start, x_end), ...]
    """
    if not words:
        return [(0, page_width)]
    
    # Agrupar palabras por su posición x inicial (redondeada a bins de 10% del ancho)
    bin_width = page_width * 0.10
    x_positions = defaultdict(int)
    
    for word in words:
        x_bin = int(word['x0'] / bin_width)
        x_positions[x_bin] += 1
    
    # Encontrar gaps significativos entre grupos de palabras
    sorted_bins = sorted(x_positions.keys())
    if len(sorted_bins) < 2:
        return [(0, page_width)]
    
    # Detectar gaps entre columnas (bins con muy pocas palabras)
    gaps = []
    for i in range(len(sorted_bins) - 1):
        current_bin = sorted_bins[i]
        next_bin = sorted_bins[i + 1]
        
        # Si hay un gap de más de 1 bin, podría ser separación de columnas
        if next_bin - current_bin > 1:
            gap_start = (current_bin + 1) * bin_width
            gap_end = next_bin * bin_width
            if (gap_end - gap_start) / page_width > COLUMN_GAP_THRESHOLD:
                gaps.append((gap_start, gap_end))
    
    # Si no hay gaps significativos, una sola columna
    if not gaps:
        return [(0, page_width)]
    
    # Construir columnas basadas en gaps
    columns = []
    prev_end = 0
    for gap_start, gap_end in gaps:
        columns.append((prev_end, gap_start))
        prev_end = gap_end
    columns.append((prev_end, page_width))
    
    # Filtrar columnas muy pequeñas
    columns = [(start, end) for start, end in columns if (end - start) / page_width > 0.1]
    
    return columns if columns else [(0, page_width)]


def _extract_text_multi_column(page) -> str:
    """
    Extrae texto de una página PDF respetando el layout multi-columna.
    
    Para CVs con diseño de dos columnas (como los de Canva), detecta las columnas
    y extrae el texto en orden lógico: primero toda la columna izquierda,
    luego toda la columna derecha.
    
    Args:
        page: Objeto página de pdfplumber
        
    Returns:
        Texto extraído en orden lógico
    """
    try:
        # Extraer palabras con coordenadas
        words = page.extract_words(
            x_tolerance=3,
            y_tolerance=3,
            keep_blank_chars=False,
            extra_attrs=['fontname', 'size']
        )
        
        if not words:
            # Fallback a extracción simple si no hay palabras
            return page.extract_text() or ""
        
        page_width = page.width
        
        # Detectar columnas
        columns = _detect_columns(words, page_width)
        
        if len(columns) <= 1:
            # Una sola columna: usar extracción estándar ordenada por y
            logger.debug("[MultiColumn] Página con 1 columna detectada")
            return page.extract_text() or ""
        
        logger.info(f"[MultiColumn] Detectadas {len(columns)} columnas: {columns}")
        
        # Agrupar palabras por columna
        column_words = [[] for _ in columns]
        for word in words:
            word_center_x = (word['x0'] + word['x1']) / 2
            for i, (col_start, col_end) in enumerate(columns):
                if col_start <= word_center_x <= col_end:
                    column_words[i].append(word)
                    break
        
        # Extraer texto de cada columna, ordenando por posición vertical
        all_text = []
        for col_idx, col_words in enumerate(column_words):
            if not col_words:
                continue
            
            # Ordenar por posición vertical (top), luego por horizontal (x0)
            col_words.sort(key=lambda w: (w['top'], w['x0']))
            
            # Reconstruir líneas agrupando palabras cercanas verticalmente
            lines = []
            current_line = []
            current_top = None
            
            for word in col_words:
                if current_top is None or abs(word['top'] - current_top) < 5:
                    current_line.append(word['text'])
                    current_top = word['top']
                else:
                    if current_line:
                        lines.append(' '.join(current_line))
                    current_line = [word['text']]
                    current_top = word['top']
            
            if current_line:
                lines.append(' '.join(current_line))
            
            column_text = '\n'.join(lines)
            if column_text.strip():
                all_text.append(column_text)
        
        # Unir columnas con separador
        return '\n\n'.join(all_text)
        
    except Exception as e:
        logger.warning(f"[MultiColumn] Error en extracción multi-columna: {e}. Usando método estándar.")
        return page.extract_text() or ""


def _extract_text_with_ocr(file_bytes: bytes) -> str:
    """
    Extrae texto de un PDF usando OCR (pytesseract + pdf2image).
    
    Este método se usa como fallback cuando pdfplumber no puede extraer texto,
    típicamente en PDFs escaneados donde el contenido son imágenes.
    
    Args:
        file_bytes: Bytes del archivo PDF
        
    Returns:
        Texto extraído via OCR
        
    Raises:
        Exception: Si el OCR falla o no está disponible
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        from PIL import Image
    except ImportError as e:
        raise Exception(
            f"OCR no disponible: {str(e)}. "
            "Instalar: pip install pytesseract pdf2image Pillow"
        )
    
    try:
        # Convertir PDF a imágenes
        logger.info("[OCR] Convirtiendo PDF a imágenes...")
        images = convert_from_bytes(file_bytes, dpi=300)
        
        if not images:
            raise Exception("No se pudieron extraer imágenes del PDF")
        
        logger.info(f"[OCR] PDF tiene {len(images)} página(s)")
        
        # Extraer texto de cada página con OCR
        all_text = []
        for i, image in enumerate(images):
            logger.info(f"[OCR] Procesando página {i+1}/{len(images)}...")
            
            # Usar español + inglés para mejor cobertura
            page_text = pytesseract.image_to_string(
                image, 
                lang='spa+eng',
                config='--psm 1'  # Automatic page segmentation with OSD
            )
            
            if page_text.strip():
                all_text.append(page_text.strip())
        
        combined_text = "\n\n".join(all_text)
        
        # Limpiar encoding
        combined_text = clean_text_encoding(combined_text)
        
        logger.info(f"[OCR] Texto extraído: {len(combined_text)} caracteres")
        
        return combined_text.strip()
        
    except Exception as e:
        error_msg = str(e)
        if "tesseract" in error_msg.lower():
            raise Exception(
                "OCR falló: Tesseract no está instalado o no se encuentra. "
                "Instalar: apt-get install tesseract-ocr tesseract-ocr-spa"
            )
        elif "poppler" in error_msg.lower() or "pdftoppm" in error_msg.lower():
            raise Exception(
                "OCR falló: Poppler no está instalado. "
                "Instalar: apt-get install poppler-utils"
            )
        else:
            raise Exception(f"OCR falló: {error_msg}")

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file with multi-column support"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    # Usar extracción multi-columna
                    page_text = _extract_text_multi_column(page)
                    if page_text:
                        text_parts.append(page_text)
                
                text = '\n\n'.join(text_parts)
                
                # Limpiar encoding solo si hay corrupción evidente
                text = clean_text_encoding(text)
                return text.strip()
        except Exception as e:
            raise Exception(f"Error extrayendo texto de PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Limpiar encoding solo si hay corrupción evidente
            text = clean_text_encoding(text)
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extrayendo texto de DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
        """Extract text from PDF bytes with multi-column support"""
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    # Usar extracción multi-columna
                    page_text = _extract_text_multi_column(page)
                    if page_text:
                        text_parts.append(page_text)
                
                text = '\n\n'.join(text_parts)
                
                # Limpiar encoding solo si hay corrupción evidente
                text = clean_text_encoding(text)
                return text.strip()
        except Exception as e:
            raise Exception(f"Error extrayendo texto de PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Limpiar encoding solo si hay corrupción evidente
            text = clean_text_encoding(text)
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extrayendo texto de DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Auto-detect file type and extract text"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return DocumentParser.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return DocumentParser.extract_text_from_docx(file_path)
        else:
            raise Exception(f"Formato de archivo no soportado: {file_ext}")
    
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from bytes based on file type.
        
        Para PDFs escaneados (sin texto extraíble), automáticamente intenta
        OCR como fallback antes de rechazar el archivo.
        """
        file_type_lower = file_type.lower()
        
        # PDF
        if file_type_lower in ['pdf', 'application/pdf']:
            text = DocumentParser.extract_text_from_pdf_bytes(file_bytes)
            
            # Verificar si el PDF tiene texto extraíble
            if not text or len(text.strip()) < MIN_TEXT_THRESHOLD:
                logger.warning(
                    f"[DocumentParser] PDF sin texto extraíble (solo {len(text.strip()) if text else 0} chars). "
                    "Intentando OCR como fallback..."
                )
                
                try:
                    # Intentar OCR como fallback
                    ocr_text = _extract_text_with_ocr(file_bytes)
                    
                    if ocr_text and len(ocr_text.strip()) >= MIN_TEXT_THRESHOLD:
                        logger.info(
                            f"[DocumentParser] OCR exitoso: {len(ocr_text)} caracteres extraídos. "
                            "CV procesado como PDF escaneado."
                        )
                        return ocr_text
                    else:
                        raise Exception(
                            "PDF escaneado: OCR no pudo extraer texto suficiente. "
                            "El archivo puede estar dañado o ser de muy baja calidad."
                        )
                        
                except Exception as ocr_error:
                    logger.error(f"[DocumentParser] OCR falló: {str(ocr_error)}")
                    raise Exception(
                        f"PDF sin texto extraíble y OCR falló: {str(ocr_error)}. "
                        "Intente convertir el documento a DOCX o a un PDF con texto seleccionable."
                    )
            
            return text
        
        # DOCX (Word moderno)
        elif file_type_lower in ['docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            return DocumentParser.extract_text_from_docx_bytes(file_bytes)
        
        # DOC (Word antiguo 97-2003) - RECHAZAR con mensaje claro
        elif file_type_lower in ['doc', 'application/msword']:
            raise Exception(
                "Formato DOC (Word 97-2003) no soportado. "
                "Por favor convierte el archivo a PDF o DOCX antes de subirlo."
            )
        
        else:
            raise Exception(f"Formato de archivo no soportado: {file_type}. Solo se permiten PDF y DOCX.")