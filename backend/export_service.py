"""
Export Service
==============
Generación de reportes PDF y DOCX para shortlists de candidatos.
Diseño ejecutivo con branding Humaniq.
"""

import os
import uuid
import logging
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional
from pathlib import Path
from io import BytesIO

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from models import User, UserRole, ExportFormat, ExportSourceType

logger = logging.getLogger(__name__)

# Paths
TEMPLATE_DIR = Path(__file__).parent / "templates"
EXPORT_DIR = Path(__file__).parent / "exports"
EXPORT_DIR.mkdir(exist_ok=True)

# Jinja2 environment
jinja_env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))


class ExportService:
    """Servicio de exportación de reportes"""
    
    def __init__(self, db, storage_service=None):
        self.db = db
        self.storage = storage_service
    
    # ========== PERMISSION CHECKS ==========
    
    def can_export(self, user: User) -> bool:
        """Verifica si el usuario puede exportar"""
        return user.role in [UserRole.SUPER_ADMIN, UserRole.ADMIN, UserRole.RECRUITER]
    
    def can_include_contact(self, user: User) -> bool:
        """Solo Admin puede incluir info de contacto"""
        return user.role in [UserRole.SUPER_ADMIN, UserRole.ADMIN]
    
    async def can_export_job(self, user: User, job_id: str) -> bool:
        """Verifica si el usuario puede exportar shortlist de una vacante"""
        if user.role in [UserRole.SUPER_ADMIN, UserRole.ADMIN]:
            return True
        
        # Recruiter puede exportar si creó la vacante
        job = await self.db.jobs.find_one({"id": job_id})
        if job and job.get("created_by") == user.id:
            return True
        
        return False
    
    # ========== DATA PREPARATION ==========
    
    async def get_job_data(self, job_id: str) -> Optional[Dict[str, Any]]:
        """Obtiene datos de una vacante"""
        job = await self.db.jobs.find_one({"id": job_id}, {"_id": 0})
        return job
    
    async def get_candidates_for_export(
        self, 
        candidate_ids: List[str],
        include_contact: bool = False
    ) -> List[Dict[str, Any]]:
        """Obtiene candidatos para exportación"""
        
        # Campos a excluir
        projection = {"_id": 0, "embedding": 0}
        if not include_contact:
            projection.update({
                "email": 0,
                "phone": 0,
                "linkedin_url": 0
            })
        
        candidates = []
        for cid in candidate_ids:
            candidate = await self.db.candidates.find_one(
                {"id": cid},
                projection
            )
            if candidate:
                candidates.append(candidate)
        
        return candidates
    
    async def get_job_matches(
        self, 
        job_id: str, 
        limit: int = 20
    ) -> List[Dict[str, Any]]:
        """Obtiene matches de una vacante con scores"""
        from job_matching_service import JobMatchingService
        from embedding_service import embedding_service
        
        # Primero obtener el job
        job = await self.db.jobs.find_one({"id": job_id}, {"_id": 0})
        if not job:
            return []
        
        matching_service = JobMatchingService(self.db, embedding_service)
        result = await matching_service.match_candidates(job, threshold=0, limit=limit)
        
        return result.get("results", [])
    
    def prepare_candidate_for_export(
        self, 
        candidate: Dict[str, Any],
        match_data: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Prepara datos del candidato para el template"""
        
        export_data = {
            "full_name": candidate.get("full_name", "Sin nombre"),
            "current_title": candidate.get("current_title"),
            "current_company": candidate.get("current_company"),
            "ai_summary": candidate.get("ai_summary"),
            "skills": candidate.get("skills", []),
            "languages": candidate.get("languages", []),
            "years_experience": candidate.get("years_experience"),
            "industry": candidate.get("industry"),
            "functional_area": candidate.get("functional_area"),
            "seniority": candidate.get("seniority"),
            "previous_companies": candidate.get("previous_companies", []),
            "education": candidate.get("education", []),
            # Contact info (si está incluido)
            "email": candidate.get("email"),
            "phone": candidate.get("phone"),
            "linkedin_url": candidate.get("linkedin_url"),
        }
        
        # Agregar datos de match si existen
        if match_data:
            export_data["match_percentage"] = match_data.get("match_percentage")
            export_data["strengths"] = match_data.get("strengths", [])
            export_data["risks"] = match_data.get("risks", [])
            export_data["missing_skills"] = match_data.get("missing_skills", [])
        
        return export_data
    
    # ========== PDF GENERATION ==========
    
    async def generate_pdf(
        self,
        job_title: str,
        candidates: List[Dict[str, Any]],
        user: User,
        client_name: Optional[str] = None,
        job_summary: Optional[str] = None,
        include_risks: bool = True,
        include_contact_info: bool = False
    ) -> bytes:
        """Genera PDF con template Jinja2 + WeasyPrint"""
        
        template = jinja_env.get_template("shortlist_report.html")
        
        html_content = template.render(
            report_title=f"Shortlist - {job_title}",
            job_title=job_title,
            client_name=client_name,
            job_summary=job_summary,
            generation_date=datetime.now().strftime("%d de %B de %Y"),
            prepared_by=user.name,
            candidate_count=len(candidates),
            candidates=candidates,
            include_risks=include_risks,
            include_contact_info=include_contact_info
        )
        
        # Generar PDF
        pdf_bytes = HTML(string=html_content).write_pdf()
        
        return pdf_bytes
    
    # ========== DOCX GENERATION ==========
    
    async def generate_docx(
        self,
        job_title: str,
        candidates: List[Dict[str, Any]],
        user: User,
        client_name: Optional[str] = None,
        job_summary: Optional[str] = None,
        include_risks: bool = True,
        include_contact_info: bool = False
    ) -> bytes:
        """Genera DOCX con python-docx"""
        
        doc = Document()
        
        # Estilos
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # === COVER PAGE ===
        # Logo/Title
        title = doc.add_paragraph()
        title_run = title.add_run("HUMANIQ")
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(26, 54, 93)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run("Executive Search")
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(44, 82, 130)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Job Title
        doc.add_paragraph("SHORTLIST DE CANDIDATOS").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        job_para = doc.add_paragraph()
        job_run = job_para.add_run(job_title)
        job_run.bold = True
        job_run.font.size = Pt(18)
        job_run.font.color.rgb = RGBColor(26, 54, 93)
        job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if client_name:
            client_para = doc.add_paragraph()
            client_para.add_run(client_name).font.size = Pt(14)
            client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        if job_summary:
            summary_para = doc.add_paragraph()
            summary_para.add_run("Perfil Buscado: ").bold = True
            summary_para.add_run(job_summary)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Meta info
        meta = doc.add_paragraph()
        meta.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")
        meta.add_run(f"Preparado por: {user.name}\n")
        meta.add_run(f"Candidatos: {len(candidates)}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # === CANDIDATE PAGES ===
        for i, candidate in enumerate(candidates):
            if i > 0:
                doc.add_page_break()
            
            # Header
            header = doc.add_paragraph()
            name_run = header.add_run(candidate.get("full_name", "Sin nombre"))
            name_run.bold = True
            name_run.font.size = Pt(16)
            name_run.font.color.rgb = RGBColor(26, 54, 93)
            
            current = doc.add_paragraph()
            title_text = candidate.get("current_title", "")
            company_text = candidate.get("current_company", "")
            if title_text:
                current.add_run(title_text)
            if company_text:
                current.add_run(f" @ {company_text}")
            
            if candidate.get("match_percentage"):
                match_para = doc.add_paragraph()
                match_run = match_para.add_run(f"Match: {candidate['match_percentage']}%")
                match_run.bold = True
                match_run.font.color.rgb = RGBColor(49, 130, 206)
            
            doc.add_paragraph()
            
            # Summary
            if candidate.get("ai_summary"):
                doc.add_paragraph().add_run("RESUMEN EJECUTIVO").bold = True
                doc.add_paragraph(candidate["ai_summary"])
                doc.add_paragraph()
            
            # Strengths
            if candidate.get("strengths"):
                doc.add_paragraph().add_run("FORTALEZAS PRINCIPALES").bold = True
                for strength in candidate["strengths"][:5]:
                    doc.add_paragraph(f"• {strength}")
                doc.add_paragraph()
            
            # Risks
            if include_risks and candidate.get("risks"):
                doc.add_paragraph().add_run("PUNTOS DE ATENCIÓN").bold = True
                for risk in candidate["risks"][:5]:
                    risk_text = risk.get("description", risk) if isinstance(risk, dict) else str(risk)
                    doc.add_paragraph(f"• {risk_text}")
                doc.add_paragraph()
            
            # Experience
            if candidate.get("previous_companies"):
                doc.add_paragraph().add_run("EXPERIENCIA PROFESIONAL").bold = True
                for exp in candidate["previous_companies"][:5]:
                    exp_para = doc.add_paragraph()
                    exp_para.add_run(f"{exp.get('title', '')}").bold = True
                    exp_para.add_run(f" - {exp.get('company_name', '')}")
                    dates = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Presente')}"
                    doc.add_paragraph(dates)
                doc.add_paragraph()
            
            # Skills
            if candidate.get("skills"):
                doc.add_paragraph().add_run("COMPETENCIAS CLAVE").bold = True
                skills_text = ", ".join(candidate["skills"][:12])
                doc.add_paragraph(skills_text)
                doc.add_paragraph()
            
            # Contact
            if include_contact_info:
                if candidate.get("email") or candidate.get("phone"):
                    doc.add_paragraph().add_run("INFORMACIÓN DE CONTACTO").bold = True
                    if candidate.get("email"):
                        doc.add_paragraph(f"Email: {candidate['email']}")
                    if candidate.get("phone"):
                        doc.add_paragraph(f"Teléfono: {candidate['phone']}")
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    # ========== MAIN EXPORT FUNCTION ==========
    
    async def export_job_shortlist(
        self,
        job_id: str,
        user: User,
        format: ExportFormat = ExportFormat.PDF,
        limit: int = 20,
        include_risks: bool = True,
        include_contact_info: bool = False,
        client_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Exporta shortlist de una vacante.
        Retorna info del archivo generado.
        """
        
        # Validar permisos
        if not self.can_export(user):
            raise PermissionError("No tienes permisos para exportar")
        
        if not await self.can_export_job(user, job_id):
            raise PermissionError("No tienes permisos para exportar esta vacante")
        
        # Validar contacto
        if include_contact_info and not self.can_include_contact(user):
            include_contact_info = False
            logger.warning(f"User {user.email} tried to include contact info without permission")
        
        # Obtener datos de la vacante
        job = await self.get_job_data(job_id)
        if not job:
            raise ValueError("Vacante no encontrada")
        
        # Obtener matches
        matches = await self.get_job_matches(job_id, limit=min(limit, 20))
        
        if not matches:
            raise ValueError("No hay candidatos para esta vacante")
        
        # Obtener datos completos de candidatos
        candidate_ids = [m["candidate_id"] for m in matches]
        candidates_data = await self.get_candidates_for_export(
            candidate_ids, 
            include_contact=include_contact_info
        )
        
        # Combinar con datos de match
        candidates_export = []
        for match in matches:
            cid = match["candidate_id"]
            candidate = next((c for c in candidates_data if c.get("id") == cid), None)
            if candidate:
                export_data = self.prepare_candidate_for_export(candidate, match)
                candidates_export.append(export_data)
        
        # Generar documento
        job_title = job.get("title", "Vacante")
        job_summary = job.get("description") or job.get("responsibilities")
        
        if format == ExportFormat.PDF:
            file_bytes = await self.generate_pdf(
                job_title=job_title,
                candidates=candidates_export,
                user=user,
                client_name=client_name or job.get("company"),
                job_summary=job_summary,
                include_risks=include_risks,
                include_contact_info=include_contact_info
            )
            file_ext = "pdf"
            content_type = "application/pdf"
        else:
            file_bytes = await self.generate_docx(
                job_title=job_title,
                candidates=candidates_export,
                user=user,
                client_name=client_name or job.get("company"),
                job_summary=job_summary,
                include_risks=include_risks,
                include_contact_info=include_contact_info
            )
            file_ext = "docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Guardar archivo
        export_id = str(uuid.uuid4())
        filename = f"shortlist_{job_id}_{export_id[:8]}.{file_ext}"
        file_path = EXPORT_DIR / filename
        
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        
        # Registrar en BD
        export_record = {
            "id": export_id,
            "user_id": user.id,
            "user_name": user.name,
            "source_type": "job",
            "source_id": job_id,
            "source_name": job_title,
            "format": format.value,
            "candidate_count": len(candidates_export),
            "candidate_ids": candidate_ids,
            "included_contact_info": include_contact_info,
            "file_path": str(file_path),
            "file_url": f"/api/exports/{export_id}/download",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await self.db.exports.insert_one(export_record)
        
        logger.info(f"Export {export_id} created by {user.email}: {len(candidates_export)} candidates, format={format.value}")
        
        return {
            "id": export_id,
            "filename": filename,
            "format": format.value,
            "candidate_count": len(candidates_export),
            "download_url": f"/api/exports/{export_id}/download",
            "included_contact_info": include_contact_info
        }
    
    async def export_custom_candidates(
        self,
        candidate_ids: List[str],
        user: User,
        title: str = "Selección de Candidatos",
        format: ExportFormat = ExportFormat.PDF,
        include_risks: bool = True,
        include_contact_info: bool = False,
        client_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Exporta una selección custom de candidatos.
        """
        
        if not self.can_export(user):
            raise PermissionError("No tienes permisos para exportar")
        
        if len(candidate_ids) > 20:
            raise ValueError("Máximo 20 candidatos por exportación")
        
        if include_contact_info and not self.can_include_contact(user):
            include_contact_info = False
        
        # Obtener candidatos
        candidates_data = await self.get_candidates_for_export(
            candidate_ids,
            include_contact=include_contact_info
        )
        
        if not candidates_data:
            raise ValueError("No se encontraron candidatos")
        
        # Preparar para exportación
        candidates_export = [
            self.prepare_candidate_for_export(c) 
            for c in candidates_data
        ]
        
        # Generar documento
        if format == ExportFormat.PDF:
            file_bytes = await self.generate_pdf(
                job_title=title,
                candidates=candidates_export,
                user=user,
                client_name=client_name,
                include_risks=include_risks,
                include_contact_info=include_contact_info
            )
            file_ext = "pdf"
        else:
            file_bytes = await self.generate_docx(
                job_title=title,
                candidates=candidates_export,
                user=user,
                client_name=client_name,
                include_risks=include_risks,
                include_contact_info=include_contact_info
            )
            file_ext = "docx"
        
        # Guardar
        export_id = str(uuid.uuid4())
        filename = f"candidates_{export_id[:8]}.{file_ext}"
        file_path = EXPORT_DIR / filename
        
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        
        # Registrar
        export_record = {
            "id": export_id,
            "user_id": user.id,
            "user_name": user.name,
            "source_type": "custom",
            "source_id": None,
            "source_name": title,
            "format": format.value,
            "candidate_count": len(candidates_export),
            "candidate_ids": candidate_ids,
            "included_contact_info": include_contact_info,
            "file_path": str(file_path),
            "file_url": f"/api/exports/{export_id}/download",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await self.db.exports.insert_one(export_record)
        
        return {
            "id": export_id,
            "filename": filename,
            "format": format.value,
            "candidate_count": len(candidates_export),
            "download_url": f"/api/exports/{export_id}/download",
            "included_contact_info": include_contact_info
        }
    
    async def get_export_record(self, export_id: str) -> Optional[Dict[str, Any]]:
        """Obtiene registro de exportación"""
        record = await self.db.exports.find_one({"id": export_id}, {"_id": 0})
        return record
    
    async def get_export_file(self, export_id: str) -> Optional[tuple]:
        """Obtiene archivo de exportación (bytes, filename, content_type)"""
        record = await self.get_export_record(export_id)
        if not record:
            return None
        
        file_path = Path(record.get("file_path", ""))
        if not file_path.exists():
            return None
        
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        
        filename = file_path.name
        
        if record.get("format") == "pdf":
            content_type = "application/pdf"
        else:
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return file_bytes, filename, content_type
    
    async def list_exports(
        self, 
        user: User,
        limit: int = 50
    ) -> List[Dict[str, Any]]:
        """Lista exportaciones del usuario (o todas si es admin)"""
        
        query = {}
        if user.role not in [UserRole.SUPER_ADMIN, UserRole.ADMIN]:
            query["user_id"] = user.id
        
        exports = await self.db.exports.find(
            query,
            {"_id": 0, "file_path": 0}
        ).sort("created_at", -1).limit(limit).to_list(limit)
        
        return exports


def create_export_service(db, storage_service=None):
    """Factory function"""
    return ExportService(db, storage_service)
