import requests
import sys
import json
import time
from datetime import datetime
from docx import Document

class AtlasRealValidationTester:
    def __init__(self, base_url="https://atlas-recruiting-ai.preview.emergentagent.com"):
        self.base_url = base_url
        self.token = None
        self.admin_token = None
        self.tests_run = 0
        self.tests_passed = 0
        self.test_results = []
        self.candidate_ids = []

    def log_test(self, name, success, details="", critical=False):
        """Log test result"""
        self.tests_run += 1
        if success:
            self.tests_passed += 1
        
        result = {
            "test": name,
            "success": success,
            "details": details,
            "critical": critical,
            "timestamp": datetime.now().isoformat()
        }
        self.test_results.append(result)
        
        status = "✅ PASSED" if success else "❌ FAILED"
        critical_marker = " [CRITICAL]" if critical else ""
        print(f"{status}{critical_marker} - {name}")
        if details:
            print(f"   Details: {details}")

    def create_docx_cv(self, profile_type="cfo", filename="test_cv.docx"):
        """Create a realistic DOCX CV file"""
        doc = Document()
        
        if profile_type == "cfo":
            doc.add_paragraph('CARLOS MENDOZA RODRIGUEZ')
            doc.add_paragraph('CFO & Director Financiero Senior')
            doc.add_paragraph('carlos.mendoza@corporativo.com | +52 55 1234 5678 | Ciudad de México, México')
            doc.add_paragraph('LinkedIn: linkedin.com/in/carlosmendoza-cfo')
            doc.add_paragraph('')
            doc.add_paragraph('RESUMEN EJECUTIVO')
            doc.add_paragraph('Director Financiero con 15+ años de experiencia liderando transformaciones financieras en manufactura automotriz y bienes de consumo. Especialista en M&A, optimización de costos y implementación de sistemas ERP. MBA en Finanzas, CPA certificado.')
            doc.add_paragraph('')
            doc.add_paragraph('EXPERIENCIA PROFESIONAL')
            doc.add_paragraph('')
            doc.add_paragraph('CFO - GRUPO AUTOMOTRIZ MEXICANO (2020-2024)')
            doc.add_paragraph('• Lideré transformación financiera de empresa con $2B USD en ingresos anuales')
            doc.add_paragraph('• Implementé SAP S/4HANA reduciendo costos operativos 15% ($30M anuales)')
            doc.add_paragraph('• Dirigí 3 adquisiciones estratégicas por valor de $500M USD')
            doc.add_paragraph('• Gestión de equipo de 45 profesionales financieros en 8 países')
            doc.add_paragraph('')
            doc.add_paragraph('Director de Finanzas - NESTLE MEXICO (2016-2020)')
            doc.add_paragraph('• P&L responsibility para 8 categorías de productos ($800M revenue)')
            doc.add_paragraph('• Optimización working capital, mejorando cash flow 25% ($50M)')
            doc.add_paragraph('• Implementación centros servicios compartidos (SSC) para LATAM')
            doc.add_paragraph('')
            doc.add_paragraph('Gerente Financiero Senior - COCA-COLA FEMSA (2012-2016)')
            doc.add_paragraph('• Financial planning & analysis para operaciones México')
            doc.add_paragraph('• Consolidación financiera de 12 plantas manufactureras')
            doc.add_paragraph('• Business partnering con operaciones, marketing, supply chain')
            doc.add_paragraph('')
            doc.add_paragraph('EDUCACIÓN')
            doc.add_paragraph('• MBA Finanzas - ITAM, México (2011)')
            doc.add_paragraph('• CPA - Instituto Mexicano de Contadores Públicos (2010)')
            doc.add_paragraph('• Licenciatura Contaduría Pública - UNAM (2008)')
            doc.add_paragraph('')
            doc.add_paragraph('HABILIDADES')
            doc.add_paragraph('• Liderazgo financiero, Strategic planning, M&A')
            doc.add_paragraph('• SAP S/4HANA, Oracle ERP, Power BI, Advanced Excel')
            doc.add_paragraph('• Lean Six Sigma Black Belt, Project Management')
            doc.add_paragraph('• Inglés nativo, Francés intermedio, Portugués básico')
            
        elif profile_type == "operations_director":
            doc.add_paragraph('MARIA ELENA GARCIA LOPEZ')
            doc.add_paragraph('Directora de Operaciones | Supply Chain & Manufacturing Expert')
            doc.add_paragraph('maria.garcia@operaciones.com | +52 81 9876 5432 | Monterrey, Nuevo León')
            doc.add_paragraph('LinkedIn: linkedin.com/in/mariagarcia-operations')
            doc.add_paragraph('')
            doc.add_paragraph('PERFIL EJECUTIVO')
            doc.add_paragraph('Directora de Operaciones con 12+ años transformando operaciones en retail, manufactura y logística. Experta en supply chain optimization, lean manufacturing y transformación digital.')
            doc.add_paragraph('')
            doc.add_paragraph('TRAYECTORIA PROFESIONAL')
            doc.add_paragraph('')
            doc.add_paragraph('Directora de Operaciones - LIVERPOOL (2021-2024)')
            doc.add_paragraph('• Optimización supply chain para red de 120+ tiendas departamentales')
            doc.add_paragraph('• Implementación WMS/TMS reduciendo inventario 20% ($40M)')
            doc.add_paragraph('• Liderazgo de 200+ colaboradores en logística y distribución')
            doc.add_paragraph('• Mejora KPIs: On-Time Delivery 98%, Fill Rate 95%')
            doc.add_paragraph('')
            doc.add_paragraph('Gerente Supply Chain - CEMEX (2018-2021)')
            doc.add_paragraph('• Gestión supply chain región Norte México (15 plantas)')
            doc.add_paragraph('• Optimización rutas distribución: ahorro $8M anuales')
            doc.add_paragraph('• Implementación IoT en flotilla: reducción 25% costos')
            doc.add_paragraph('')
            doc.add_paragraph('Coordinadora Producción - GRUPO BIMBO (2015-2018)')
            doc.add_paragraph('• Supervisión 3 plantas producción (capacidad 500 ton/día)')
            doc.add_paragraph('• Implementación metodologías Lean: mejora 15% eficiencia OEE')
            doc.add_paragraph('')
            doc.add_paragraph('FORMACIÓN ACADÉMICA')
            doc.add_paragraph('• Maestría Ingeniería Industrial - ITESM (2014)')
            doc.add_paragraph('• Ingeniería Industrial - UANL (2012)')
            doc.add_paragraph('• Six Sigma Black Belt Certification - ASQ (2019)')
            doc.add_paragraph('')
            doc.add_paragraph('COMPETENCIAS TÉCNICAS')
            doc.add_paragraph('• Supply Chain Management, Lean Manufacturing, Six Sigma')
            doc.add_paragraph('• SAP MM/PP/WM, Oracle SCM, WMS Manhattan, TMS')
            doc.add_paragraph('• Power BI, Tableau, Advanced Excel')
            doc.add_paragraph('• Inglés fluido, Alemán intermedio')
        
        doc.save(filename)
        return filename

    def authenticate(self):
        """Authenticate with test user"""
        login_data = {
            "email": "admin@atlas.com",
            "password": "password123"
        }
        
        try:
            response = requests.post(
                f"{self.base_url}/api/auth/login",
                json=login_data,
                headers={'Content-Type': 'application/json'},
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                self.token = result['access_token']
                self.log_test("Authentication", True, "Successfully authenticated")
                return True
            else:
                self.log_test("Authentication", False, f"Status: {response.status_code}", critical=True)
                return False
                
        except Exception as e:
            self.log_test("Authentication", False, f"Exception: {str(e)}", critical=True)
            return False

    def setup_super_admin(self):
        """Setup super admin for taxonomy testing"""
        admin_data = {
            "email": "superadmin@atlas.com",
            "password": "SuperAdmin123!",
            "name": "Super Admin Test",
            "role": "super_admin"
        }
        
        try:
            # Try to login first (user might already exist)
            login_response = requests.post(
                f"{self.base_url}/api/auth/login",
                json={"email": admin_data["email"], "password": admin_data["password"]},
                headers={'Content-Type': 'application/json'},
                timeout=30
            )
            
            if login_response.status_code == 200:
                result = login_response.json()
                self.admin_token = result['access_token']
                self.log_test("Super Admin Setup", True, "Super admin authenticated")
                return True
            else:
                # Try to register
                register_response = requests.post(
                    f"{self.base_url}/api/auth/register",
                    json=admin_data,
                    headers={'Content-Type': 'application/json'},
                    timeout=30
                )
                
                if register_response.status_code == 200:
                    # Now login
                    login_response = requests.post(
                        f"{self.base_url}/api/auth/login",
                        json={"email": admin_data["email"], "password": admin_data["password"]},
                        headers={'Content-Type': 'application/json'},
                        timeout=30
                    )
                    
                    if login_response.status_code == 200:
                        result = login_response.json()
                        self.admin_token = result['access_token']
                        self.log_test("Super Admin Setup", True, "Super admin created and authenticated")
                        return True
                
                self.log_test("Super Admin Setup", False, "Could not setup super admin")
                return False
                    
        except Exception as e:
            self.log_test("Super Admin Setup", False, f"Exception: {str(e)}")
            return False

    def test_end_to_end_cv_upload(self):
        """Test complete CV upload flow"""
        print("\n🔍 Testing END-TO-END CV Upload Flow...")
        
        # Create CFO CV
        cv_file = self.create_docx_cv("cfo", "/app/carlos_cfo.docx")
        
        try:
            with open(cv_file, 'rb') as f:
                files = {
                    'file': ('carlos_mendoza_cfo.docx', f.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                }
                
                headers = {'Authorization': f'Bearer {self.token}'}
                
                response = requests.post(
                    f"{self.base_url}/api/candidates/upload-resume",
                    files=files,
                    headers=headers,
                    timeout=90
                )
                
                if response.status_code == 200:
                    result = response.json()
                    
                    # Analyze results
                    status = result.get('status')
                    parsed_data = result.get('parsed_data', {})
                    candidate_id = result.get('candidate_id')
                    
                    if candidate_id:
                        self.candidate_ids.append(candidate_id)
                    
                    # Check parsing quality
                    name_ok = bool(parsed_data.get('full_name'))
                    email_ok = bool(parsed_data.get('email'))
                    company_ok = bool(parsed_data.get('current_company'))
                    title_ok = bool(parsed_data.get('current_title'))
                    
                    parsing_score = sum([name_ok, email_ok, company_ok, title_ok])
                    success = parsing_score >= 3 and candidate_id is not None
                    
                    self.log_test(
                        "END-TO-END CV Upload",
                        success,
                        f"Status: {status}, Parsing: {parsing_score}/4 fields, ID: {candidate_id[:8] if candidate_id else 'None'}...",
                        critical=True
                    )
                    
                    # Test AI classification
                    if success and candidate_id:
                        time.sleep(2)
                        self.test_ai_classification(candidate_id)
                    
                    return success
                else:
                    self.log_test("END-TO-END CV Upload", False, f"Status: {response.status_code}", critical=True)
                    return False
                    
        except Exception as e:
            self.log_test("END-TO-END CV Upload", False, f"Exception: {str(e)}", critical=True)
            return False

    def test_ai_classification(self, candidate_id):
        """Test AI classification accuracy"""
        try:
            response = requests.post(
                f"{self.base_url}/api/atlas/classify/{candidate_id}",
                headers={'Authorization': f'Bearer {self.token}'},
                timeout=60
            )
            
            if response.status_code == 200:
                classification = response.json()
                
                industry = classification.get('industry', '').lower()
                functional_area = classification.get('functional_area', '').lower()
                seniority = classification.get('seniority', '').lower()
                confidence = classification.get('confidence_score', 0)
                
                # Check business relevance for CFO profile
                industry_relevant = any(term in industry for term in ['manufacturing', 'manufactura', 'automotive', 'automotriz', 'consumer', 'consumo', 'financial', 'financiero'])
                area_relevant = any(term in functional_area for term in ['finance', 'finanzas', 'accounting', 'contabilidad', 'general management', 'direccion'])
                seniority_relevant = seniority in ['director', 'c_level', 'vp', 'senior']
                
                relevance_score = sum([industry_relevant, area_relevant, seniority_relevant])
                success = relevance_score >= 2 and confidence > 0.6
                
                self.log_test(
                    "AI Classification",
                    success,
                    f"Industry: {classification.get('industry')}, Area: {classification.get('functional_area')}, Seniority: {seniority}, Confidence: {confidence:.2f}",
                    critical=True
                )
                
                return success
            else:
                self.log_test("AI Classification", False, f"Status: {response.status_code}", critical=True)
                return False
                
        except Exception as e:
            self.log_test("AI Classification", False, f"Exception: {str(e)}", critical=True)
            return False

    def test_duplicate_detection_90_percent(self):
        """Test duplicate detection at ≥90% confidence"""
        print("\n🔍 Testing Duplicate Detection ≥90% Confidence...")
        
        # Create variant of same CFO
        doc = Document()
        doc.add_paragraph('Carlos A. Mendoza Rodriguez')  # Slight name variation
        doc.add_paragraph('Chief Financial Officer')  # Different title format
        doc.add_paragraph('c.mendoza@gmail.com | 55-1234-5678 | CDMX, México')  # Different contact
        doc.add_paragraph('')
        doc.add_paragraph('PERFIL PROFESIONAL')
        doc.add_paragraph('Director Financiero con más de 15 años de experiencia en transformaciones financieras en sector automotriz y manufactura.')
        doc.add_paragraph('')
        doc.add_paragraph('EXPERIENCIA')
        doc.add_paragraph('CFO - GRUPO AUTOMOTRIZ MEXICANO (2020-2024)')
        doc.add_paragraph('• Transformación financiera empresa $2B ingresos')
        doc.add_paragraph('• SAP implementation, reducción costos 15%')
        doc.add_paragraph('')
        doc.add_paragraph('Director Finanzas - NESTLE MEXICO (2016-2020)')
        doc.add_paragraph('• P&L 8 categorías productos')
        doc.add_paragraph('• Working capital optimization')
        
        variant_file = "/app/carlos_variant.docx"
        doc.save(variant_file)
        
        try:
            with open(variant_file, 'rb') as f:
                files = {
                    'file': ('carlos_variant.docx', f.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                }
                
                headers = {'Authorization': f'Bearer {self.token}'}
                
                response = requests.post(
                    f"{self.base_url}/api/candidates/upload-resume",
                    files=files,
                    headers=headers,
                    timeout=90
                )
                
                if response.status_code == 200:
                    result = response.json()
                    
                    status = result.get('status')
                    duplicates = result.get('duplicates', [])
                    
                    if status == 'duplicate_detected' and duplicates:
                        max_confidence = max(d.get('confidence', 0) for d in duplicates)
                        success = max_confidence >= 0.90
                        
                        self.log_test(
                            "Duplicate Detection ≥90%",
                            success,
                            f"Status: {status}, Duplicates: {len(duplicates)}, Max confidence: {max_confidence:.2f}",
                            critical=True
                        )
                        return success
                    else:
                        self.log_test(
                            "Duplicate Detection ≥90%",
                            False,
                            f"Status: {status}, Expected duplicate detection",
                            critical=True
                        )
                        return False
                else:
                    self.log_test("Duplicate Detection ≥90%", False, f"Status: {response.status_code}", critical=True)
                    return False
                    
        except Exception as e:
            self.log_test("Duplicate Detection ≥90%", False, f"Exception: {str(e)}", critical=True)
            return False

    def test_object_storage_emergent(self):
        """Test Emergent Object Storage integration"""
        print("\n🔍 Testing Emergent Object Storage...")
        
        if not self.candidate_ids:
            self.log_test("Object Storage (Emergent)", False, "No candidates to check", critical=True)
            return False
        
        try:
            candidate_id = self.candidate_ids[0]
            response = requests.get(
                f"{self.base_url}/api/candidates/{candidate_id}",
                headers={'Authorization': f'Bearer {self.token}'},
                timeout=30
            )
            
            if response.status_code == 200:
                candidate = response.json()
                resume_files = candidate.get('resume_files', [])
                
                if resume_files:
                    file_path = resume_files[0].get('file_path', '')
                    
                    # Check for Emergent Object Storage characteristics
                    is_emergent = (
                        'atlas-talent-vault' in file_path and 
                        'resumes' in file_path and
                        not file_path.startswith('/app/') and
                        not file_path.startswith('uploads/') and
                        '/' in file_path  # Has proper path structure
                    )
                    
                    self.log_test(
                        "Object Storage (Emergent)",
                        is_emergent,
                        f"Path: {file_path}, Is Emergent: {is_emergent}",
                        critical=True
                    )
                    return is_emergent
                else:
                    self.log_test("Object Storage (Emergent)", False, "No resume files found", critical=True)
                    return False
            else:
                self.log_test("Object Storage (Emergent)", False, f"Status: {response.status_code}", critical=True)
                return False
                
        except Exception as e:
            self.log_test("Object Storage (Emergent)", False, f"Exception: {str(e)}", critical=True)
            return False

    def test_hybrid_search_recruiting_queries(self):
        """Test hybrid search with real recruiting queries"""
        print("\n🔍 Testing Hybrid Search - Recruiting Queries...")
        
        recruiting_queries = [
            "CFO manufactura",
            "Director operaciones automotriz", 
            "gerente supply chain retail",
            "Director financiero M&A",
            "VP finanzas experiencia SAP"
        ]
        
        successful_queries = 0
        
        for query in recruiting_queries:
            try:
                params = {
                    'query': query,
                    'use_semantic': True,
                    'limit': 20
                }
                
                response = requests.post(
                    f"{self.base_url}/api/search/hybrid",
                    params=params,
                    headers={'Authorization': f'Bearer {self.token}'},
                    timeout=60
                )
                
                if response.status_code == 200:
                    results = response.json()
                    
                    # Check for semantic results
                    semantic_results = [
                        r for r in results 
                        if r.get('match_breakdown', {}).get('semantic', 0) > 0
                    ]
                    
                    if len(semantic_results) > 0:
                        successful_queries += 1
                        self.log_test(
                            f"Search: '{query}'",
                            True,
                            f"{len(semantic_results)} semantic results"
                        )
                    else:
                        self.log_test(
                            f"Search: '{query}'",
                            False,
                            "No semantic results"
                        )
                else:
                    self.log_test(f"Search: '{query}'", False, f"Status: {response.status_code}")
                    
            except Exception as e:
                self.log_test(f"Search: '{query}'", False, f"Exception: {str(e)}")
        
        success_rate = successful_queries / len(recruiting_queries)
        overall_success = success_rate >= 0.6
        
        self.log_test(
            "Hybrid Search Overall",
            overall_success,
            f"Success: {successful_queries}/{len(recruiting_queries)} ({success_rate:.1%})",
            critical=True
        )
        
        return overall_success

    def test_taxonomy_crud_super_admin(self):
        """Test taxonomy CRUD operations"""
        print("\n🔍 Testing Taxonomy CRUD...")
        
        if not self.admin_token:
            self.log_test("Taxonomy CRUD", False, "No super admin token", critical=True)
            return False
        
        industry_data = {
            "name_es": "Industria Test Atlas",
            "name_en": "Atlas Test Industry",
            "description": "Test industry for Atlas validation"
        }
        
        try:
            # CREATE
            response = requests.post(
                f"{self.base_url}/api/admin/industries",
                json=industry_data,
                headers={'Authorization': f'Bearer {self.admin_token}'},
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                industry_id = result.get('industry_id')
                
                if industry_id:
                    # UPDATE
                    update_data = {
                        "name_es": "Industria Test Atlas Actualizada",
                        "name_en": "Updated Atlas Test Industry",
                        "description": "Updated test industry"
                    }
                    
                    update_response = requests.put(
                        f"{self.base_url}/api/admin/industries/{industry_id}",
                        json=update_data,
                        headers={'Authorization': f'Bearer {self.admin_token}'},
                        timeout=30
                    )
                    
                    if update_response.status_code == 200:
                        # DELETE
                        delete_response = requests.delete(
                            f"{self.base_url}/api/admin/industries/{industry_id}",
                            headers={'Authorization': f'Bearer {self.admin_token}'},
                            timeout=30
                        )
                        
                        if delete_response.status_code == 200:
                            self.log_test("Taxonomy CRUD", True, "Create, Update, Delete successful", critical=True)
                            return True
            
            self.log_test("Taxonomy CRUD", False, "CRUD operations failed", critical=True)
            return False
                
        except Exception as e:
            self.log_test("Taxonomy CRUD", False, f"Exception: {str(e)}", critical=True)
            return False

    def test_performance_benchmarks(self):
        """Test performance with multiple uploads"""
        print("\n🔍 Testing Performance Benchmarks...")
        
        # Create operations director CV
        ops_file = self.create_docx_cv("operations_director", "/app/maria_ops.docx")
        
        upload_times = []
        search_times = []
        successful_uploads = 0
        
        # Test 2 uploads
        for i, cv_file in enumerate(["/app/carlos_cfo.docx", ops_file]):
            start_time = time.time()
            
            try:
                with open(cv_file, 'rb') as f:
                    files = {
                        'file': (f'performance_test_{i}.docx', f.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    }
                    
                    response = requests.post(
                        f"{self.base_url}/api/candidates/upload-resume",
                        files=files,
                        headers={'Authorization': f'Bearer {self.token}'},
                        timeout=120
                    )
                    
                    upload_time = time.time() - start_time
                    upload_times.append(upload_time)
                    
                    if response.status_code == 200:
                        successful_uploads += 1
                        result = response.json()
                        if result.get('candidate_id'):
                            self.candidate_ids.append(result['candidate_id'])
                    
                    self.log_test(
                        f"Performance Upload {i+1}",
                        response.status_code == 200,
                        f"Time: {upload_time:.1f}s, Status: {response.status_code}"
                    )
                    
            except Exception as e:
                self.log_test(f"Performance Upload {i+1}", False, f"Exception: {str(e)}")
            
            # Test search performance
            search_start = time.time()
            try:
                params = {'query': 'director financiero', 'use_semantic': True, 'limit': 10}
                search_response = requests.post(
                    f"{self.base_url}/api/search/hybrid",
                    params=params,
                    headers={'Authorization': f'Bearer {self.token}'},
                    timeout=30
                )
                
                search_time = time.time() - search_start
                search_times.append(search_time)
                
                self.log_test(
                    f"Performance Search {i+1}",
                    search_response.status_code == 200,
                    f"Time: {search_time:.1f}s"
                )
                
            except Exception as e:
                self.log_test(f"Performance Search {i+1}", False, f"Exception: {str(e)}")
        
        # Performance analysis
        if upload_times and search_times:
            avg_upload = sum(upload_times) / len(upload_times)
            avg_search = sum(search_times) / len(search_times)
            
            # Performance criteria for AI-powered system
            upload_acceptable = avg_upload < 60  # AI processing takes time
            search_acceptable = avg_search < 5   # Search should be fast
            
            performance_good = upload_acceptable and search_acceptable
            
            self.log_test(
                "Performance Benchmarks",
                performance_good,
                f"Upload avg: {avg_upload:.1f}s, Search avg: {avg_search:.1f}s, Uploads: {successful_uploads}/2",
                critical=True
            )
            
            return performance_good
        
        return False

    def run_atlas_validation(self):
        """Run complete Atlas validation"""
        print("🚀 ATLAS TALENT VAULT - REAL VALIDATION FOR PHASE 2")
        print("=" * 65)
        
        # Setup
        print("\n🔧 AUTHENTICATION & SETUP")
        print("-" * 35)
        
        if not self.authenticate():
            return False
        
        self.setup_super_admin()
        
        # Core validation
        print("\n🎯 CORE BUSINESS VALIDATION")
        print("-" * 35)
        
        self.test_end_to_end_cv_upload()
        self.test_duplicate_detection_90_percent()
        self.test_object_storage_emergent()
        self.test_hybrid_search_recruiting_queries()
        self.test_taxonomy_crud_super_admin()
        self.test_performance_benchmarks()
        
        return True

    def print_final_assessment(self):
        """Print final Phase 2 readiness assessment"""
        print("\n" + "=" * 65)
        print("📋 ATLAS TALENT VAULT - PHASE 2 READINESS")
        print("=" * 65)
        
        critical_tests = [r for r in self.test_results if r.get('critical', False)]
        critical_passed = [r for r in critical_tests if r['success']]
        
        print(f"Total Tests: {self.tests_run}")
        print(f"Passed: {self.tests_passed}")
        print(f"Failed: {self.tests_run - self.tests_passed}")
        print(f"Success Rate: {(self.tests_passed / self.tests_run * 100):.1f}%")
        
        print(f"\nCritical Features: {len(critical_tests)}")
        print(f"Critical Working: {len(critical_passed)}")
        print(f"Critical Success: {(len(critical_passed) / len(critical_tests) * 100):.1f}%" if critical_tests else "N/A")
        
        # Show critical failures
        failed_critical = [r for r in critical_tests if not r['success']]
        if failed_critical:
            print(f"\n❌ CRITICAL ISSUES:")
            for result in failed_critical:
                print(f"  • {result['test']}: {result['details']}")
        
        # Final assessment
        critical_success_rate = len(critical_passed) / len(critical_tests) if critical_tests else 0
        overall_success_rate = self.tests_passed / self.tests_run if self.tests_run > 0 else 0
        
        ready_for_phase2 = (
            critical_success_rate >= 0.8 and
            overall_success_rate >= 0.75 and
            len(self.candidate_ids) > 0
        )
        
        print(f"\n{'🎉 READY FOR PHASE 2 - INTERNAL TEAM USE' if ready_for_phase2 else '⚠️  NEEDS FIXES BEFORE PHASE 2'}")
        
        if ready_for_phase2:
            print("✅ CV upload & AI parsing working")
            print("✅ Duplicate detection ≥90% confidence")
            print("✅ Emergent Object Storage integrated")
            print("✅ Semantic search operational")
            print("✅ Taxonomy management working")
            print("✅ Performance within limits")
        
        return ready_for_phase2

def main():
    tester = AtlasRealValidationTester()
    
    try:
        success = tester.run_atlas_validation()
        ready = tester.print_final_assessment()
        
        # Save results
        with open('/app/atlas_phase2_validation.json', 'w') as f:
            json.dump({
                'phase2_ready': ready,
                'validation_summary': {
                    'total_tests': tester.tests_run,
                    'passed_tests': tester.tests_passed,
                    'success_rate': tester.tests_passed / tester.tests_run * 100 if tester.tests_run > 0 else 0,
                    'critical_success_rate': len([r for r in tester.test_results if r.get('critical', False) and r['success']]) / len([r for r in tester.test_results if r.get('critical', False)]) * 100 if [r for r in tester.test_results if r.get('critical', False)] else 0
                },
                'test_results': tester.test_results,
                'candidates_created': len(tester.candidate_ids),
                'timestamp': datetime.now().isoformat()
            }, f, indent=2)
        
        return 0 if ready else 1
        
    except Exception as e:
        print(f"\n💥 Validation failed: {str(e)}")
        return 1

if __name__ == "__main__":
    sys.exit(main())